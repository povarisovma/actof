from docx.shared import Inches, Pt
import docx
from pytils import dt
import os
import win32com.client as com
import shutil
import re
import settings
import datetime


def textforlist(textinput):
    textlst = []
    for line in textinput:
        textlst.append(line.strip())
    return textlst


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def get_from_bodylist_azsnum(blist):
    for i in range(len(blist[0].split())):
        if 'АЗС' in blist[0].split()[i]:
            if re.sub("\\D", "", blist[0].split()[i + 1]).isdigit():
                return re.sub("\\D", "", blist[0].split()[i + 1])
            elif re.sub("\\D", "", blist[0].split()[i + 2]).isdigit():
                return re.sub("\\D", "", blist[0].split()[i + 2])


def get_from_bodylist_ssonum(blist):
    for i in range(len(blist[0].split())):
        if 'ССО' in blist[0].split()[i]:
            if re.sub("\\D", "", blist[0].split()[i + 1]).isdigit():
                return re.sub("\\D", "", blist[0].split()[i + 1])
            elif re.sub("\\D", "", blist[0].split()[i + 2]).isdigit():
                return re.sub("\\D", "", blist[0].split()[i + 2])


def del_empty_paragraphs(doc, btext):
    empty_parag = 45 - 3 - len(btext) - 2
    for p in doc.paragraphs:
        if p.text == '' and empty_parag != 0:
            delete_paragraph(p)
            empty_parag -= 1


def get_current_date():
    return str(dt.ru_strftime("%d %B %Y" + ' г.', inflected=True))


def get_number_act():
    filelist = os.listdir(settings.get_general_acts_path_folder())
    numacts = []
    nextnumact = ''
    for i in range(len(filelist)):
        if '.docx' in filelist[i]:
            numacts.append(int(filelist[i].split('_')[0][3:]))
    numacts.sort(reverse=True)
    for i in range(len(numacts)):
        # print(i, numacts[i])
        if 0 <= (numacts[i] - numacts[i + 1]) <= 2:
            nextnumact = str(numacts[i] + 1)
            break
    return nextnumact


def get_listdir_pdf_files_in_dict(path):
    generallist = []
    filelist = os.listdir(path)
    if filelist:
        pdflist = []
        for i in filelist:
            if '.pdf' in i:
                pdflist.append(i)
        for file in pdflist:
            generallist.append({"title": file, "creating": datetime.datetime.fromtimestamp(os.path.getctime(os.path.join(path, file))),
                               "modifine": datetime.datetime.fromtimestamp(os.path.getmtime(os.path.join(path, file)))})
    return generallist


def get_listdir_docx_files_in_dict(path):
    generallist = []
    filelist = os.listdir(path)
    if filelist:
        docxlist = []
        for i in filelist:
            if '.docx' in i:
                docxlist.append(i)
        for file in docxlist:
            generallist.append({"title": file, "creating": datetime.datetime.fromtimestamp(os.path.getctime(os.path.join(path, file))),
                               "modifine": datetime.datetime.fromtimestamp(os.path.getmtime(os.path.join(path, file)))})
    return generallist


def create_docx_file_from_bodylist(blist, AZSnum, SSOnum, ACTnum, docxpath):
    doc = docx.Document(settings.get_docx_templ_path())
    nowdate = get_current_date()
    bottext = False
    numline = 0

    for i in range(len(doc.paragraphs)):
        if i == 0:
            hd = doc.paragraphs[i]
            hd.paragraph_format.space_after = Pt(10)
            doc.paragraphs[i].text = 'АКТ № ' + ACTnum
            doc.paragraphs[i].style = 'Normal'
            doc.paragraphs[i].alignment = 1
            doc.paragraphs[i].runs[0].bold = True
            doc.paragraphs[i].runs[0].font.name = 'Times New Roman'
            doc.paragraphs[i].runs[0].font.size = Pt(14)
            continue
        if i == 1:
            hd2 = doc.paragraphs[i]
            hd2.paragraph_format.space_after = Pt(10)
            doc.paragraphs[i].style = 'Normal'
            doc.paragraphs[i].add_run('АЗС №' + AZSnum + '	ССО №' + SSOnum + '\t\t\t\t\t\t\t\t' + nowdate)
            doc.paragraphs[i].runs[0].font.name = 'Times New Roman'
            doc.paragraphs[i].runs[0].font.size = Pt(12)
            continue
        if i > 1 and not bottext and numline < len(blist):
            p2 = doc.paragraphs[i]
            run2 = p2.add_run('\t' + blist[numline])
            p2.paragraph_format.alignment = 3
            if blist[numline].find('В связи с чем') != -1:
                bottext = True
                p2.paragraph_format.space_before = Pt(0)
                p2.paragraph_format.space_after = Pt(0)
            else:
                p2.paragraph_format.space_before = Pt(0)
                p2.paragraph_format.space_after = Pt(0)
            p2.paragraph_format.line_spacing = Pt(0)
            run2.font.name = 'Times New Roman'
            run2.font.size = Pt(12)
            numline += 1
            continue
        if i > 1 and bottext and numline < len(blist):
            p3 = doc.paragraphs[i]
            doc.paragraphs[i].style = 'List Paragraph'
            run3 = p3.add_run('— ' + blist[numline])
            run3.font.name = 'Times New Roman'
            run3.font.size = Pt(12)
            p3.paragraph_format.left_indent = Inches(0.8)
            p3.paragraph_format.alignment = 0
            p3.paragraph_format.line_spacing = Pt(0)
            p3.paragraph_format.space_before = Pt(0)
            p3.paragraph_format.space_after = Pt(0)
            numline += 1
            continue
    del_empty_paragraphs(doc, blist)
    doc.save(docxpath)


def create_pdf_file_from_docx(docxpath):
    wdFormatPDF = 17
    out_file = docxpath.strip(".docx")
    word = com.DispatchEx('word.application')
    doccon = word.Documents.Open(docxpath)
    doccon.SaveAs(out_file, FileFormat=wdFormatPDF)
    doccon.Close()
    word.Quit()


def get_path_to_file_to_string(filename):
    return settings.get_local_acts_path_folder() + filename


def get_name_pdf_from_docx(filenamedocx):
    return filenamedocx.strip(".docx") + ".pdf"


def copy_files_to_general_folder(filenamedocx):
    docxpath = get_path_to_file_to_string(filenamedocx)
    filenamepdf = get_name_pdf_from_docx(filenamedocx)
    docxpathgen = settings.get_general_acts_path_folder() + filenamedocx
    pdfpath = settings.get_local_acts_path_folder() + filenamepdf
    pdfpathgen = settings.get_general_acts_path_folder() + filenamepdf
    shutil.copyfile(docxpath, docxpathgen)
    shutil.copyfile(pdfpath, pdfpathgen)


def create_docx_and_pdf_files(lst):
    AZSnum = get_from_bodylist_azsnum(lst)
    SSOnum = get_from_bodylist_ssonum(lst)
    ACTnum = get_number_act()
    filenamedocx = f"Акт{ACTnum}_АЗС{AZSnum}_ССО{SSOnum}.docx"
    docxpath = settings.get_local_acts_path_folder() + filenamedocx
    create_docx_file_from_bodylist(lst, AZSnum, SSOnum, ACTnum, docxpath)
    create_pdf_file_from_docx(docxpath)
    copy_files_to_general_folder(filenamedocx)
