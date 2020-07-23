from docx.shared import Mm, Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT, WD_TAB_LEADER
import docx


text = 'Настоящим подтверждаю, что 8.05.2020 на АЗС №31052 в результате сбоя, пролив топлива АИ-92 объемом 30,12 л.' \
       ' на сумму 1249,98 руб. был учтен в ССО №58 как пролив объемом 12,21 л. на сумму 506,72 руб. из-за чего' \
       ' возникло расхождение по счетчикам ТРК. Оплата производилась по банковской карте №****4240 по RRN №' \
       ' 012923616562.'
text2 = 'Также подтверждаю, что 8.05.2020 на АЗС №30052 в результате сбоя, пролив топлива АИ-92 объемом 12,21 л. на' \
        ' сумму 506 руб. не был учтен в ССО №58. Оплата производилась за наличный расчет.'
text3 = 'В связи с чем, прошу: '
text4 = 'Считать пролив топлива АИ-92 по банковской карте №****4240 от 8.05.2020 объемом 30,12 л. на сумму 1249,98 руб.'
text5 = 'Пролив топлива АИ-92 объемом 12,21 л. на сумму 506 руб. считать проливом за наличный расчет.'
text6 = 'Ведущий специалист'
text7 = 'Отдел поддержки СУ НАМОС    	                                   	_______________Поварисов М.А.'
text71 = 'Отдел поддержки СУ НАМОС'
text72 = 'Поварисов М.А.'


#Разметка страницы:
document = docx.Document()
section = document.sections[0]
#Высота листа:
section.page_height = Mm(297)
#Ширина листа:
section.page_width = Mm(210)
#Отступ слева:
section.left_margin = Mm(23.3)
#Отступ справа:
section.right_margin = Mm(11.1)
#Отступ сверху:
section.top_margin = Mm(30.6)
#Отступ снизу:
section.bottom_margin = Mm(3)
#Отступ от верхнего колонтитула:
section.header_distance = Mm(15.9)
#Отступ от нижнего колонтитула:
section.footer_distance = Mm(1.3)

#Добавление заголовка:
document.add_heading('АКТ № 98587')
#Установка стиля заголовка:
document.paragraphs[0].style = 'Normal'
#Установка жирного текста:
document.paragraphs[0].runs[0].bold = True
#Установка шрифта заголовка:
document.paragraphs[0].runs[0].font.name = 'Times New Roman'
#Установка размера шрифта:
document.paragraphs[0].runs[0].font.size = Pt(14)
#Установка выравнивания заголовка по центру:
document.paragraphs[0].alignment = 1

#Добавление параграфа с шапкой номер АЗС ССО и Дата. Используется 8 табуляций.
document.add_paragraph('АЗС №31052	ССО №58' + '\t\t\t\t\t\t\t\t' + '17 мая 2020 г.', 'Normal')
#Установка шрифта:
document.paragraphs[1].runs[0].font.name = 'Times New Roman'
#Установка размера шрифта:
document.paragraphs[1].runs[0].font.size = Pt(12)

#Создание переменной для редактирования нового параграфа:
p1 = document.add_paragraph()
#Создание переменной для редактирования текста:
run1 = p1.add_run('\t' + text)
#Установка выравнивания по ширине:
p1.paragraph_format.alignment = 3
#Отступ между строк внутри параграфа:
p1.paragraph_format.line_spacing = Pt(0)
#Отступ от параграфа:
p1.paragraph_format.space_before = Pt(0)
#Отступ после параграфа:
p1.paragraph_format.space_after = Pt(0)
#Установка шрифта и его размера:
run1.font.name = 'Times New Roman'
run1.font.size = Pt(12)

p2 = document.add_paragraph()
run2 = p2.add_run('\t' + text2)
p2.paragraph_format.alignment = 3
p2.paragraph_format.line_spacing = Pt(0)
p2.paragraph_format.space_before = Pt(0)
p2.paragraph_format.space_after = Pt(0)
run2.font.name = 'Times New Roman'
run2.font.size = Pt(12)

pcentreline = document.add_paragraph()
runcentreline = pcentreline.add_run('\t' + text3)
pcentreline.paragraph_format.alignment = 3
pcentreline.paragraph_format.line_spacing = Pt(0)
pcentreline.paragraph_format.space_before = Pt(0)
pcentreline.paragraph_format.space_after = Pt(0)
runcentreline.font.name = 'Times New Roman'
runcentreline.font.size = Pt(12)

p3 = document.add_paragraph()
p3.style = 'List Bullet 3'
run3 = p3.add_run(text4)
p3.paragraph_format.alignment = 3
p3.paragraph_format.line_spacing = Pt(0)
p3.paragraph_format.space_before = Pt(0)
p3.paragraph_format.space_after = Pt(0)
run3.font.name = 'Times New Roman'
run3.font.size = Pt(12)

p4 = document.add_paragraph()
p4.style = 'List Bullet 3'
run4 = p4.add_run(text5)
p4.paragraph_format.alignment = 3
p4.paragraph_format.line_spacing = Pt(0)
p4.paragraph_format.space_before = Pt(0)
p4.paragraph_format.space_after = Pt(0)
run4.font.name = 'Times New Roman'
run4.font.size = Pt(12)

psign = document.add_paragraph()
runsign = psign.add_run(text6)
# psign.paragraph_format.alignment = 1
psign.paragraph_format.line_spacing = Pt(0)
psign.paragraph_format.space_before = Pt(20)
psign.paragraph_format.space_after = Pt(0)
runsign.font.name = 'Times New Roman'
runsign.font.bold = True
runsign.font.size = Pt(12)

psign1 = document.add_paragraph()
runsign1 = psign1.add_run(text71)
runsign1.add_text('\t\t\t\t\t\t\t')
# runsign1.add_picture('sign.jpg')
runsign1.add_text(text72)
psign1.paragraph_format.line_spacing = Pt(0)
psign1.paragraph_format.space_before = Pt(0)
psign1.paragraph_format.space_after = Pt(0)
runsign1.font.name = 'Times New Roman'
runsign1.font.bold = True
runsign1.font.size = Pt(12)

psign2 = document.add_paragraph()
runsign2 = psign2.add_run()
runsign2.add_text('\t\t\t\t\t\t\t\t')
runsign2.add_picture('sign.jpg')
psign2.paragraph_format.line_spacing = Pt(0)
psign2.paragraph_format.space_before = Pt(0)
psign2.paragraph_format.space_after = Pt(0)
runsign2.font.name = 'Times New Roman'
runsign2.font.bold = True
runsign2.font.size = Pt(12)

# document.add_picture('sign.jpg')

#---------------Старый код, оставил для примера:
# document.add_paragraph('\t' + text, 'Normal')
# for i in range(len(document.paragraphs[2].runs)):
#     document.paragraphs[2].runs[i].font.name = 'Times New Roman'
#     document.paragraphs[2].runs[i].font.size = Pt(12)
# document.paragraphs[2].alignment = 3
#
# document.add_paragraph('\t' + text2, 'Normal')
# for i in range(len(document.paragraphs[2].runs)):
#     document.paragraphs[3].runs[i].font.name = 'Times New Roman'
#     document.paragraphs[3].runs[i].font.size = Pt(12)
# document.paragraphs[3].alignment = 3
#---------------Конец старого кода!

document.save('demo.docx')