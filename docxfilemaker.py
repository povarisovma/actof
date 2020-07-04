from docx.shared import Inches, Pt
import docx
import getnumberact
import datetime
from pytils import dt
import os
import win32com.client as com
import shutil
import re


def textforlist(textinput):
    textlst = []
    for line in textinput:
        textlst.append(line.strip())
    return textlst


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def createdocxnpdffiles(lst):
    # print(os.path.exists('local_acts'))
    if not os.path.exists('local_acts'):
        os.mkdir('local_acts')


    # filein = open('input.txt', 'r', encoding='utf-8')
    doc = docx.Document('template.docx')


    para_list = []
    for par in range(len(doc.paragraphs)):
        para_list.append(doc.paragraphs[par])


    bottext = False
    # textlist = textforlist(filein)
    # filein.close()
    textlist = lst
    AZSnum = ''
    SSOnum = ''
    ACTnum = getnumberact.get_number_act()
    for i in range(len(textlist[0].split())):
        if 'АЗС' in textlist[0].split()[i]:
            AZSnum = re.sub("\\D", "", textlist[0].split()[i + 1])
        if 'ССО' in textlist[0].split()[i]:
            SSOnum = re.sub("\\D", "", textlist[0].split()[i + 1])
    now = datetime.datetime.now()
    nowdate = str(dt.ru_strftime("%d %B %Y" + ' г.', inflected=True))
    numline = 0
    for i in range(len(para_list)):
        if i == 0:
            hd = para_list[i]
            hd.paragraph_format.space_after = Pt(10)
            para_list[i].text = 'АКТ № ' + ACTnum
            para_list[i].style = 'Normal'
            para_list[i].alignment = 1
            para_list[i].runs[0].bold = True
            para_list[i].runs[0].font.name = 'Times New Roman'
            para_list[i].runs[0].font.size = Pt(14)
            continue
        if i == 1:
            hd2 = para_list[i]
            hd2.paragraph_format.space_after = Pt(10)
            para_list[i].style = 'Normal'
            para_list[i].add_run('АЗС №' + AZSnum + '	ССО №' + SSOnum + '\t\t\t\t\t\t\t\t' + nowdate)
            para_list[i].runs[0].font.name = 'Times New Roman'
            para_list[i].runs[0].font.size = Pt(12)
            continue
        if i > 1 and not bottext and numline < len(textlist):
            p2 = para_list[i]
            run2 = p2.add_run('\t' + textlist[numline])

            p2.paragraph_format.alignment = 3
            if textlist[numline].find('В связи с чем') != -1:
                bottext = True
                p2.paragraph_format.space_before = Pt(0)
                p2.paragraph_format.space_after = Pt(0)
            else:
                p2.paragraph_format.space_before = Pt(0)
                p2.paragraph_format.space_after = Pt(0)
            p2.paragraph_format.line_spacing = Pt(0)
            run2.font.name = 'Times New Roman'
            run2.font.size = Pt(12)
            numline += 1
            continue
        if i > 1 and bottext and numline < len(textlist):
            p3 = para_list[i]
            para_list[i].style = 'List Paragraph'
            run3 = p3.add_run('— ' + textlist[numline])
            run3.font.name = 'Times New Roman'
            run3.font.size = Pt(12)
            p3.paragraph_format.left_indent = Inches(0.8)
            p3.paragraph_format.alignment = 0
            p3.paragraph_format.line_spacing = Pt(0)
            p3.paragraph_format.space_before = Pt(0)
            p3.paragraph_format.space_after = Pt(0)
            numline += 1
            continue

    empty_parag = 45 - 3 - len(textlist) - 2
    for p in doc.paragraphs:
        if p.text == '' and empty_parag != 0:
            delete_paragraph(p)
            empty_parag -= 1

    filenamedocx = 'local_acts\\' + 'Акт' + ACTnum + '_' + 'АЗС' + AZSnum + '_' + 'ССО' + SSOnum + '.docx'
    filenamepdf = 'local_acts\\' + 'Акт' + ACTnum + '_' + 'АЗС' + AZSnum + '_' + 'ССО' + SSOnum + '.pdf'
    endfiledocx = getnumberact.get_path() + '\\' + 'Акт' + ACTnum + '_' + 'АЗС' + AZSnum + '_' + 'ССО' + SSOnum + '.docx'
    endfilepdf = getnumberact.get_path() + '\\' + 'Акт' + ACTnum + '_' + 'АЗС' + AZSnum + '_' + 'ССО' + SSOnum + '.pdf'
    # print(filenamedocx, filenamepdf, sep='\n')

    doc.save(filenamedocx)

    wdFormatPDF = 17

    in_file = os.path.abspath(filenamedocx)
    out_file = os.path.abspath(filenamepdf)
    # print(in_file)
    # print(out_file)

    word = com.DispatchEx('word.application')
    doccon = word.Documents.Open(in_file)
    doccon.SaveAs(out_file, FileFormat=wdFormatPDF)
    doccon.Close()
    word.Quit()


    shutil.copyfile(filenamedocx, endfiledocx)
    shutil.copyfile(filenamepdf, endfilepdf)
    # print(ACTnum, AZSnum, SSOnum)
    return 'Акт' + ACTnum + '_' + 'АЗС' + AZSnum + '_' + 'ССО' + SSOnum + '.pdf'
